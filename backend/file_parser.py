from io import BytesIO
from typing import Dict, List, Tuple, Any, Optional
import re
import logging

import fitz  # PyMuPDF
import pdfplumber
import docx
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)


class LegalDocumentParser:
    """
    Robust legal document parser:
      - PDF: text + structured tables + image detection + OCR fallback
      - DOCX: paragraphs + structured tables + image count
      - TXT: plain text
      - Extracts legal elements via regex
    """

    def __init__(
        self,
        ocr_if_no_text: bool = True,
        ocr_min_chars_threshold: int = 40,
        ocr_zoom: float = 2.0,
    ):
        """
        :param ocr_if_no_text: Run OCR when a page has little/no selectable text.
        :param ocr_min_chars_threshold: If page text length < this and page has images,
                                        attempt OCR.
        :param ocr_zoom: Scaling factor for rasterization before OCR for better accuracy.
        """
        self.ocr_if_no_text = ocr_if_no_text
        self.ocr_min_chars_threshold = ocr_min_chars_threshold
        self.ocr_zoom = ocr_zoom

        # Pre-compile regexes for speed & maintainability
        self.legal_patterns = {
            "contract_clauses": [
                r"WHEREAS\s+.*?;",
                r"NOW,?\s+THEREFORE\s+.*?;",
                r"IN\s+WITNESS\s+WHEREOF\s+.*?;",
            ],
            "legal_citations": [
                r"\b\d+\s+[A-Z][a-z]+\.?\s+\d+\b",   # Case cites: 410 U.S. 113 (genericized)
                r"\b\d+\s+U\.S\.C\.?\s+§?\s*\d+[a-zA-Z0-9\-]*\b",  # 15 U.S.C. § 78j(b)
                r"\b\d+\s+C\.F\.R\.?\s+§?\s*\d+(?:\.\d+)*\b",      # 12 C.F.R. § 1026
            ],
            "dates": [
                r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
                r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b",
            ],
        }
        # Compile
        self._compiled_patterns: Dict[str, List[re.Pattern]] = {
            k: [re.compile(p, re.IGNORECASE | re.MULTILINE) for p in pats]
            for k, pats in self.legal_patterns.items()
        }

    # ----------------------------- Public API -----------------------------

    def extract_text(self, filename: str, file_bytes: bytes) -> Dict[str, Any]:
        """Main entry point with type dispatch."""
        name = filename.lower()
        if name.endswith(".pdf"):
            return self.extract_text_from_pdf(file_bytes)
        elif name.endswith(".docx"):
            return self.extract_text_from_docx(file_bytes)
        elif name.endswith(".txt"):
            text = file_bytes.decode("utf-8", errors="ignore").strip()
            return {
                "text": text,
                "metadata": {
                    "source_type": "txt",
                    "legal_elements": self._extract_legal_elements(text),
                },
            }
        else:
            raise ValueError(f"Unsupported file type: {filename}")

    def extract_text_from_pdf(self, file_bytes: bytes) -> Dict[str, Any]:
        """
        Enhanced PDF extraction:
          - pdfplumber text & tables
          - fallback to PyMuPDF text if needed
          - detect images; OCR if page text is sparse
          - structured metadata (tables per page, images per page, OCR flags)
        """
        full_text_parts: List[str] = []
        tables_all: List[Dict[str, Any]] = []
        images_all: List[Dict[str, Any]] = []
        ocr_pages: List[int] = []

        pages_count = 0

        # Attempt pdfplumber for text & tables first
        plumber_ok = False
        try:
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                pages_count = len(pdf.pages)
                plumber_ok = True
                # Open a parallel PyMuPDF doc for image detection and as a fallback source
                with fitz.open(stream=file_bytes, filetype="pdf") as fdoc:
                    for page_num, (pl_page, fz_page) in enumerate(zip(pdf.pages, fdoc), start=1):
                        page_text = (pl_page.extract_text() or "").strip()

                        # Extract structured tables with pdfplumber
                        page_tables = []
                        try:
                            raw_tables = pl_page.extract_tables()
                            if raw_tables:
                                for t in raw_tables:
                                    if not t:
                                        continue
                                    rows = [[cell if cell is not None else "" for cell in row] for row in t]
                                    page_tables.append(rows)
                                    # Also append a readable preview into the full text
                                    preview = "\n".join(["\t".join(r) for r in rows])
                                    full_text_parts.append(f"\n[TABLE p{page_num}]\n{preview}\n[/TABLE]\n")
                        except Exception as e:
                            logger.debug(f"pdfplumber table extraction failed on page {page_num}: {e}")

                        # Image info from PyMuPDF
                        img_infos = self._get_fitz_image_info(fz_page)
                        if img_infos:
                            images_all.extend(
                                {"page": page_num, **info} for info in img_infos
                            )

                        # If no/low text and there are images, consider OCR
                        if self.ocr_if_no_text and (not page_text or len(page_text) < self.ocr_min_chars_threshold):
                            if img_infos:  # likely scanned
                                ocr_text = self._ocr_page(fz_page, zoom=self.ocr_zoom)
                                if ocr_text.strip():
                                    page_text = (page_text + "\n" + ocr_text).strip() if page_text else ocr_text.strip()
                                    ocr_pages.append(page_num)

                        # If still empty, fallback to fitz text
                        if not page_text:
                            fz_text = (fz_page.get_text() or "").strip()
                            if fz_text:
                                page_text = fz_text

                        if page_text:
                            full_text_parts.append(f"\n--- Page {page_num} ---\n{page_text}\n")

                        if page_tables:
                            tables_all.append({"page": page_num, "rows": page_tables})
        except Exception as e:
            logger.debug(f"pdfplumber failed entirely, falling back to PyMuPDF: {e}")

        # If pdfplumber failed completely, do a fitz-only pass
        if not plumber_ok:
            with fitz.open(stream=file_bytes, filetype="pdf") as fdoc:
                pages_count = len(fdoc)
                for page_index, fz_page in enumerate(fdoc, start=1):
                    page_text = (fz_page.get_text() or "").strip()

                    # Image info
                    img_infos = self._get_fitz_image_info(fz_page)
                    if img_infos:
                        images_all.extend({"page": page_index, **info} for info in img_infos)

                    # OCR decision
                    if self.ocr_if_no_text and (not page_text or len(page_text) < self.ocr_min_chars_threshold):
                        if img_infos:
                            ocr_text = self._ocr_page(fz_page, zoom=self.ocr_zoom)
                            if ocr_text.strip():
                                page_text = (page_text + "\n" + ocr_text).strip() if page_text else ocr_text.strip()
                                ocr_pages.append(page_index)

                    if page_text:
                        full_text_parts.append(f"\n--- Page {page_index} ---\n{page_text}\n")

            # Note: no table extraction in this branch (plumber unavailable)

        full_text = "".join(full_text_parts).strip()

        metadata = {
            "source_type": "pdf",
            "pages": pages_count,
            "tables_extracted": len(tables_all),
            "images_detected": len(images_all),
            "ocr_pages": ocr_pages,  # which pages needed OCR
            "tables": tables_all,    # structured tables: [{"page": n, "rows":[ [..], .. ]}]
            "images": images_all,    # images metadata per page (bbox, width, height)
            "legal_elements": self._extract_legal_elements(full_text),
        }

        return {"text": full_text, "metadata": metadata}

    def extract_text_from_docx(self, file_bytes: bytes) -> Dict[str, Any]:
        """DOCX: paragraphs + structured tables + image count (inline shapes)."""
        document = docx.Document(BytesIO(file_bytes))

        paragraphs: List[str] = []
        tables_all: List[Dict[str, Any]] = []

        # paragraphs
        for p in document.paragraphs:
            if p.text and p.text.strip():
                paragraphs.append(p.text.strip())

        # tables
        for t_index, table in enumerate(document.tables, start=1):
            rows_struct = []
            for row in table.rows:
                rows_struct.append([cell.text or "" for cell in row.cells])

            if rows_struct:
                tables_all.append({"table_index": t_index, "rows": rows_struct})

        # inline image count (simple signal; extracting image bytes requires deeper access)
        try:
            inline_img_count = len(document.inline_shapes)  # type: ignore[attr-defined]
        except Exception:
            inline_img_count = 0

        # Compose full text with readable table previews
        full_text_parts: List[str] = []
        if paragraphs:
            full_text_parts.append("\n".join(paragraphs))

        for t in tables_all:
            preview = "\n".join(["\t".join(r) for r in t["rows"]])
            full_text_parts.append(f"\n[TABLE docx:{t['table_index']}]\n{preview}\n[/TABLE]\n")

        full_text = "\n".join(full_text_parts).strip()

        metadata = {
            "source_type": "docx",
            "paragraphs": len(paragraphs),
            "tables_extracted": len(tables_all),
            "images_detected": inline_img_count,
            "tables": tables_all,
            "legal_elements": self._extract_legal_elements(full_text),
        }

        return {"text": full_text, "metadata": metadata}

    # ----------------------------- Helpers -----------------------------

    def _extract_legal_elements(self, text: str) -> List[Dict[str, Any]]:
        """Extract legal-specific elements from text."""
        elements: List[Dict[str, Any]] = []
        if not text:
            return elements

        for category, patterns in self._compiled_patterns.items():
            for pat in patterns:
                for match in pat.findall(text):
                    elements.append(
                        {
                            "type": category,
                            "text": match.strip(),
                            "pattern": pat.pattern,
                        }
                    )
        return elements

    def _get_fitz_image_info(self, fz_page: "fitz.Page") -> List[Dict[str, Any]]:
        """
        Return basic metadata for images on a page (no bytes extracted).
        Each entry includes bbox (rect) and image dimensions if available.
        """
        infos: List[Dict[str, Any]] = []
        try:
            # get_images(full=True) gives xref and more; we can also map to bbox via page.get_text("rawdict")
            # Here we use a layout walk to approximate image rectangles.
            raw = fz_page.get_text("rawdict")
            blocks = raw.get("blocks", []) if isinstance(raw, dict) else []
            for b in blocks:
                if b.get("type") == 1:  # image block
                    # 'bbox': [x0, y0, x1, y1]
                    bbox = b.get("bbox")
                    width = int((bbox[2] - bbox[0])) if bbox else None
                    height = int((bbox[3] - bbox[1])) if bbox else None
                    infos.append(
                        {
                            "bbox": bbox,
                            "width": width,
                            "height": height,
                        }
                    )
        except Exception as e:
            logger.debug(f"Failed to read image info: {e}")
        return infos

    def _ocr_page(self, fz_page: "fitz.Page", zoom: float = 2.0) -> str:
        """
        Rasterize a PyMuPDF page and run Tesseract OCR.
        :param zoom: scale factor (2.0 is a good default for legibility).
        """
        try:
            mat = fitz.Matrix(zoom, zoom)
            pix = fz_page.get_pixmap(matrix=mat, alpha=False)
            img = Image.open(BytesIO(pix.tobytes()))
            text = pytesseract.image_to_string(img)
            if text and text.strip():
                return text.strip()
        except Exception as e:
            logger.debug(f"OCR failed on page {fz_page.number + 1}: {e}")
        return ""


# ----------------------------- Usage Example -----------------------------
# parser = LegalDocumentParser()
# result = parser.extract_text("contract.pdf", open("contract.pdf", "rb").read())
# print(result["metadata"])
# print(result["text"][:2000])
